# gcrbot/src/gcrbot/tools_document.py
"""
Tools pour l'agent Document : lecture, embedding et recherche dans les fichiers.
Supporte : PDF, Word (.docx), Excel (.xlsx), Text (.txt)
"""

import os
import re
import json
import hashlib
import logging
from typing import Optional, List, Dict, Any
from datetime import datetime
from crewai.tools import tool

# Configuration logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("gcrbot.tools_document")

# Dossier de stockage des documents - Chemin ABSOLU depuis la racine du projet
# Structure: GCRBOT/gcrbot/src/gcrbot/tools_document.py
# Donc on remonte: tools_document.py -> gcrbot -> src -> gcrbot -> GCRBOT
_CURRENT_FILE = os.path.abspath(__file__)  # .../GCRBOT/gcrbot/src/gcrbot/tools_document.py
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(_CURRENT_FILE))))  # .../GCRBOT
DOC_DB_PATH = os.path.join(_PROJECT_ROOT, 'docDB')
EMBEDDINGS_FILE = os.path.join(DOC_DB_PATH, 'embeddings_index.json')

# Log du chemin pour debug
logger.info(f"📁 DOC_DB_PATH: {DOC_DB_PATH}")
logger.info(f"📁 PROJECT_ROOT: {_PROJECT_ROOT}")

# Imports conditionnels pour les différents formats
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("pdfplumber non installé - PDF non supporté")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx non installé - Word non supporté")

try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False
    logger.warning("openpyxl non installé - Excel non supporté")


# ═══════════════════════════════════════════════════════════════════════════════
# FONCTIONS UTILITAIRES
# ═══════════════════════════════════════════════════════════════════════════════

def get_file_hash(filepath: str) -> str:
    """Génère un hash unique pour un fichier."""
    with open(filepath, 'rb') as f:
        return hashlib.md5(f.read()).hexdigest()[:12]


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Découpe le texte en chunks avec overlap pour meilleur contexte."""
    chunks = []
    start = 0
    text = text.strip()
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Essayer de couper à une fin de phrase
        if end < len(text):
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            cut_point = max(last_period, last_newline)
            if cut_point > chunk_size // 2:
                chunk = chunk[:cut_point + 1]
                end = start + cut_point + 1
        
        if chunk.strip():
            chunks.append(chunk.strip())
        
        start = end - overlap
        if start < 0:
            start = end
    
    return chunks


def simple_embedding(text: str) -> List[float]:
    """
    Crée un embedding simple basé sur les mots-clés.
    Pour une vraie production, utiliser Gemini ou autre modèle.
    """
    # Normaliser le texte
    text_lower = text.lower()
    words = re.findall(r'\b\w+\b', text_lower)
    
    # Créer un vecteur de caractéristiques simples
    features = []
    
    # Fréquence de mots courants
    common_words = ['le', 'la', 'de', 'et', 'en', 'un', 'une', 'est', 'pour', 'que',
                    'the', 'a', 'an', 'is', 'for', 'that', 'with', 'on', 'at', 'to']
    
    for word in common_words:
        features.append(words.count(word) / max(len(words), 1))
    
    # Longueur normalisée
    features.append(min(len(text) / 5000, 1.0))
    features.append(len(words) / 500)
    
    # Présence de chiffres
    features.append(len(re.findall(r'\d+', text)) / 100)
    
    return features


def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """Calcule la similarité cosinus entre deux vecteurs."""
    if len(vec1) != len(vec2) or not vec1:
        return 0.0
    
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    magnitude1 = sum(a ** 2 for a in vec1) ** 0.5
    magnitude2 = sum(b ** 2 for b in vec2) ** 0.5
    
    if magnitude1 == 0 or magnitude2 == 0:
        return 0.0
    
    return dot_product / (magnitude1 * magnitude2)


# ═══════════════════════════════════════════════════════════════════════════════
# EXTRACTION DE CONTENU
# ═══════════════════════════════════════════════════════════════════════════════

def extract_pdf_text(filepath: str) -> str:
    """Extrait le texte d'un fichier PDF."""
    if not PDF_SUPPORT:
        return "Erreur: pdfplumber non installé"
    
    try:
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"[Page {i}]\n{page_text}")
                
                # Extraire aussi les tableaux
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        for row in table:
                            if row:
                                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                if row_text.strip():
                                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Erreur extraction PDF: {e}")
        return f"Erreur extraction PDF: {str(e)}"


def extract_docx_text(filepath: str) -> str:
    """Extrait le texte d'un fichier Word (.docx)."""
    if not DOCX_SUPPORT:
        return "Erreur: python-docx non installé"
    
    try:
        doc = DocxDocument(filepath)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extraire aussi les tableaux
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Erreur extraction DOCX: {e}")
        return f"Erreur extraction DOCX: {str(e)}"


def extract_excel_text(filepath: str) -> str:
    """Extrait le texte d'un fichier Excel (.xlsx)."""
    if not EXCEL_SUPPORT:
        return "Erreur: openpyxl non installé"
    
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Feuille: {sheet_name} ===")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    text_parts.append(" | ".join(row_values))
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Erreur extraction Excel: {e}")
        return f"Erreur extraction Excel: {str(e)}"


def extract_txt_text(filepath: str) -> str:
    """Extrait le texte d'un fichier texte."""
    try:
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        return "Erreur: encodage non reconnu"
    except Exception as e:
        logger.error(f"Erreur lecture TXT: {e}")
        return f"Erreur lecture TXT: {str(e)}"


def extract_document_content(filepath: str) -> str:
    """Extrait le contenu d'un document selon son type."""
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.pdf':
        return extract_pdf_text(filepath)
    elif ext == '.docx':
        return extract_docx_text(filepath)
    elif ext in ['.xlsx', '.xls']:
        return extract_excel_text(filepath)
    elif ext == '.txt':
        return extract_txt_text(filepath)
    else:
        return f"Format non supporté: {ext}"


# ═══════════════════════════════════════════════════════════════════════════════
# GESTION DE L'INDEX
# ═══════════════════════════════════════════════════════════════════════════════

def load_embeddings_index() -> Dict[str, Any]:
    """Charge l'index des embeddings."""
    if os.path.exists(EMBEDDINGS_FILE):
        try:
            with open(EMBEDDINGS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return {"documents": {}, "chunks": []}


def save_embeddings_index(index: Dict[str, Any]):
    """Sauvegarde l'index des embeddings."""
    os.makedirs(DOC_DB_PATH, exist_ok=True)
    with open(EMBEDDINGS_FILE, 'w', encoding='utf-8') as f:
        json.dump(index, f, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════════════════════════════
# TOOLS CREWAI
# ═══════════════════════════════════════════════════════════════════════════════

@tool("process_uploaded_document")
def process_uploaded_document(filepath: str) -> str:
    """
    Traite un document uploadé : extraction, chunking et indexation.
    
    Args:
        filepath: Chemin complet vers le fichier uploadé
    
    Returns:
        Confirmation du traitement avec statistiques
    """
    try:
        logger.info(f"📄 Traitement document: {filepath}")
        
        if not os.path.exists(filepath):
            return f"❌ Fichier non trouvé: {filepath}"
        
        filename = os.path.basename(filepath)
        file_hash = get_file_hash(filepath)
        
        # Extraire le contenu
        content = extract_document_content(filepath)
        
        if content.startswith("Erreur"):
            return content
        
        # Découper en chunks
        chunks = chunk_text(content)
        
        if not chunks:
            return "❌ Aucun contenu extractible du document"
        
        # Créer les embeddings et indexer
        index = load_embeddings_index()
        
        # Enregistrer le document
        doc_info = {
            "filename": filename,
            "filepath": filepath,
            "hash": file_hash,
            "processed_at": datetime.now().isoformat(),
            "total_chunks": len(chunks),
            "total_chars": len(content)
        }
        index["documents"][file_hash] = doc_info
        
        # Indexer les chunks
        for i, chunk in enumerate(chunks):
            chunk_entry = {
                "doc_hash": file_hash,
                "chunk_id": i,
                "text": chunk,
                "embedding": simple_embedding(chunk)
            }
            index["chunks"].append(chunk_entry)
        
        save_embeddings_index(index)
        
        logger.info(f"✅ Document indexé: {len(chunks)} chunks")
        
        return f"""✅ Document traité avec succès !

📄 Fichier: {filename}
📊 Statistiques:
- {len(chunks)} sections indexées
- {len(content):,} caractères extraits
- Prêt pour les questions !

💡 Vous pouvez maintenant poser des questions sur ce document."""
        
    except Exception as e:
        logger.error(f"Erreur traitement document: {e}")
        return f"❌ Erreur traitement: {str(e)}"


@tool("search_in_documents")
def search_in_documents(query: str, top_k: int = 5) -> str:
    """
    Recherche dans tous les documents indexés.
    
    Args:
        query: Question ou mots-clés à rechercher
        top_k: Nombre de résultats à retourner (défaut: 5)
    
    Returns:
        Les passages les plus pertinents trouvés
    """
    try:
        logger.info(f"🔍 Recherche: {query[:50]}...")
        
        index = load_embeddings_index()
        
        if not index["chunks"]:
            return "❌ Aucun document indexé. Veuillez d'abord uploader un document."
        
        # Créer l'embedding de la requête
        query_embedding = simple_embedding(query)
        
        # Calculer les similarités
        results = []
        for chunk in index["chunks"]:
            similarity = cosine_similarity(query_embedding, chunk["embedding"])
            
            # Bonus si les mots de la requête sont présents
            query_words = set(re.findall(r'\b\w+\b', query.lower()))
            chunk_words = set(re.findall(r'\b\w+\b', chunk["text"].lower()))
            word_overlap = len(query_words & chunk_words) / max(len(query_words), 1)
            
            score = similarity * 0.4 + word_overlap * 0.6
            
            results.append({
                "score": score,
                "text": chunk["text"],
                "doc_hash": chunk["doc_hash"],
                "chunk_id": chunk["chunk_id"]
            })
        
        # Trier par score
        results.sort(key=lambda x: x["score"], reverse=True)
        top_results = results[:top_k]
        
        if not top_results or top_results[0]["score"] < 0.1:
            return "❌ Aucun passage pertinent trouvé pour cette question."
        
        # Formater les résultats
        output = ["📚 RÉSULTATS DE RECHERCHE", "=" * 50]
        
        for i, result in enumerate(top_results, 1):
            doc_info = index["documents"].get(result["doc_hash"], {})
            filename = doc_info.get("filename", "Document inconnu")
            
            output.append(f"\n📄 Résultat {i} (score: {result['score']:.2f})")
            output.append(f"Source: {filename}")
            output.append("-" * 40)
            
            # Limiter la longueur du texte affiché
            text = result["text"]
            if len(text) > 800:
                text = text[:800] + "..."
            output.append(text)
        
        return "\n".join(output)
        
    except Exception as e:
        logger.error(f"Erreur recherche: {e}")
        return f"❌ Erreur recherche: {str(e)}"


@tool("summarize_document")
def summarize_document(doc_name: str = "") -> str:
    """
    Génère un résumé d'un document ou de tous les documents.
    
    Args:
        doc_name: Nom du document (optionnel, sinon résume tous)
    
    Returns:
        Résumé du/des document(s)
    """
    try:
        index = load_embeddings_index()
        
        if not index["documents"]:
            return "❌ Aucun document indexé."
        
        # Trouver le(s) document(s) à résumer
        target_docs = []
        
        if doc_name:
            for doc_hash, doc_info in index["documents"].items():
                if doc_name.lower() in doc_info["filename"].lower():
                    target_docs.append((doc_hash, doc_info))
        else:
            target_docs = list(index["documents"].items())
        
        if not target_docs:
            return f"❌ Document '{doc_name}' non trouvé."
        
        output = ["📋 RÉSUMÉ DES DOCUMENTS", "=" * 50]
        
        for doc_hash, doc_info in target_docs:
            filename = doc_info["filename"]
            total_chunks = doc_info["total_chunks"]
            total_chars = doc_info["total_chars"]
            
            # Récupérer les premiers et derniers chunks pour le résumé
            doc_chunks = [c for c in index["chunks"] if c["doc_hash"] == doc_hash]
            
            output.append(f"\n📄 {filename}")
            output.append(f"   📊 {total_chunks} sections, {total_chars:,} caractères")
            output.append(f"   📅 Indexé le: {doc_info.get('processed_at', 'N/A')[:10]}")
            
            if doc_chunks:
                # Aperçu du début
                first_chunk = doc_chunks[0]["text"][:300]
                output.append(f"\n   📝 Aperçu:\n   {first_chunk}...")
        
        output.append("\n" + "=" * 50)
        output.append("💡 Posez des questions pour explorer le contenu en détail.")
        
        return "\n".join(output)
        
    except Exception as e:
        logger.error(f"Erreur résumé: {e}")
        return f"❌ Erreur résumé: {str(e)}"


@tool("list_documents")
def list_documents() -> str:
    """
    Liste tous les documents indexés.
    
    Returns:
        Liste des documents avec leurs informations
    """
    try:
        index = load_embeddings_index()
        
        if not index["documents"]:
            return "📭 Aucun document indexé pour le moment.\n💡 Uploadez un document pour commencer !"
        
        output = ["📚 DOCUMENTS INDEXÉS", "=" * 50]
        
        for i, (doc_hash, doc_info) in enumerate(index["documents"].items(), 1):
            filename = doc_info["filename"]
            chunks = doc_info["total_chunks"]
            chars = doc_info["total_chars"]
            date = doc_info.get("processed_at", "")[:10]
            
            output.append(f"\n{i}. 📄 {filename}")
            output.append(f"   └─ {chunks} sections | {chars:,} caractères | {date}")
        
        output.append(f"\n{'=' * 50}")
        output.append(f"📊 Total: {len(index['documents'])} document(s)")
        
        return "\n".join(output)
        
    except Exception as e:
        logger.error(f"Erreur liste: {e}")
        return f"❌ Erreur: {str(e)}"


def _search_documents_internal(query: str, top_k: int = 5) -> str:
    """Fonction interne de recherche (sans décorateur @tool)."""
    try:
        logger.info(f"🔍 Recherche interne: {query[:50]}...")
        
        index = load_embeddings_index()
        
        if not index["chunks"]:
            return "❌ Aucun document indexé. Veuillez d'abord uploader un document."
        
        # Créer l'embedding de la requête
        query_embedding = simple_embedding(query)
        
        # Calculer les similarités
        results = []
        for chunk in index["chunks"]:
            similarity = cosine_similarity(query_embedding, chunk["embedding"])
            
            # Bonus si les mots de la requête sont présents
            query_words = set(re.findall(r'\b\w+\b', query.lower()))
            chunk_words = set(re.findall(r'\b\w+\b', chunk["text"].lower()))
            word_overlap = len(query_words & chunk_words) / max(len(query_words), 1)
            
            score = similarity * 0.4 + word_overlap * 0.6
            
            results.append({
                "score": score,
                "text": chunk["text"],
                "doc_hash": chunk["doc_hash"],
                "chunk_id": chunk["chunk_id"]
            })
        
        # Trier par score
        results.sort(key=lambda x: x["score"], reverse=True)
        top_results = results[:top_k]
        
        if not top_results or top_results[0]["score"] < 0.1:
            return "❌ Aucun passage pertinent trouvé pour cette question."
        
        # Formater les résultats
        output = ["📚 RÉSULTATS DE RECHERCHE", "=" * 50]
        
        for i, result in enumerate(top_results, 1):
            doc_info = index["documents"].get(result["doc_hash"], {})
            filename = doc_info.get("filename", "Document inconnu")
            
            output.append(f"\n📄 Résultat {i} (score: {result['score']:.2f})")
            output.append(f"Source: {filename}")
            output.append("-" * 40)
            
            text = result["text"]
            if len(text) > 800:
                text = text[:800] + "..."
            output.append(text)
        
        return "\n".join(output)
        
    except Exception as e:
        logger.error(f"Erreur recherche interne: {e}")
        return f"❌ Erreur recherche: {str(e)}"


def _generate_response_with_llm(question: str, context: str, source_files: list) -> str:
    """
    Utilise Gemini pour générer une réponse intelligente basée sur le contexte.
    """
    try:
        import google.generativeai as genai
        import time
        
        # Configurer Gemini
        api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
        if not api_key:
            logger.warning("Clé API Gemini non trouvée, retour au mode simple")
            return None
        
        genai.configure(api_key=api_key)
        
        # Utiliser un modèle plus léger pour éviter les erreurs de quota
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Créer le prompt
        sources_str = ", ".join(set(source_files)) if source_files else "Document uploadé"
        
        prompt = f"""Tu es un assistant expert en analyse de documents. 
Réponds à la question de l'utilisateur en te basant UNIQUEMENT sur le contenu fourni ci-dessous.

📄 CONTENU DU DOCUMENT:
{context}

❓ QUESTION DE L'UTILISATEUR:
{question}

📝 INSTRUCTIONS:
- Réponds dans la MÊME LANGUE que la question (français si la question est en français)
- Si la question demande un RÉSUMÉ → fais un résumé clair et structuré
- Si la question demande une REFORMULATION → reformule le contenu différemment
- Si la question demande une EXPLICATION → explique en détail avec des exemples si possible
- Si la question est SPÉCIFIQUE (ex: "c'est quoi l'IoC?") → réponds précisément à cette question
- Sois concis mais complet
- Ne dis PAS "selon le document" ou "le document mentionne", réponds directement
- À la fin, cite la source: 📄 Source: {sources_str}

RÉPONSE:"""
        
        response = model.generate_content(prompt)
        
        if response and response.text:
            return response.text.strip()
        
        return None
        
    except Exception as e:
        error_str = str(e)
        if "429" in error_str or "quota" in error_str.lower() or "exhausted" in error_str.lower():
            logger.warning(f"⚠️ Quota API dépassé, utilisation du mode simple")
        else:
            logger.error(f"Erreur génération LLM: {e}")
        return None


def _simple_answer_from_context(question: str, context: str, source_files: list) -> str:
    """
    Génère une réponse simple sans LLM quand le quota est dépassé.
    Extrait les passages les plus pertinents.
    """
    sources_str = ", ".join(set(source_files)) if source_files else "Document uploadé"
    
    # Limiter le contexte
    if len(context) > 1500:
        context = context[:1500] + "..."
    
    return f"""📖 Voici les informations trouvées dans vos documents :

{context}

📄 Source: {sources_str}

💡 Note: Pour des réponses plus détaillées, réessayez dans quelques minutes (quota API temporairement dépassé)."""


@tool("answer_from_document")
def answer_from_document(question: str) -> str:
    """
    Répond à une question en utilisant le contenu des documents indexés.
    Utilise Gemini pour générer une réponse intelligente et contextuelle.
    
    Args:
        question: La question de l'utilisateur
    
    Returns:
        Réponse basée sur le contenu des documents
    """
    try:
        logger.info(f"🔍 Question RAG: {question[:50]}...")
        
        # Charger l'index
        index = load_embeddings_index()
        
        if not index["chunks"]:
            return "❌ Aucun document indexé. Veuillez d'abord uploader un document."
        
        # Rechercher les passages pertinents
        query_embedding = simple_embedding(question)
        
        results = []
        for chunk in index["chunks"]:
            similarity = cosine_similarity(query_embedding, chunk["embedding"])
            
            query_words = set(re.findall(r'\b\w+\b', question.lower()))
            chunk_words = set(re.findall(r'\b\w+\b', chunk["text"].lower()))
            word_overlap = len(query_words & chunk_words) / max(len(query_words), 1)
            
            score = similarity * 0.4 + word_overlap * 0.6
            
            results.append({
                "score": score,
                "text": chunk["text"],
                "doc_hash": chunk["doc_hash"]
            })
        
        # Trier et prendre les meilleurs
        results.sort(key=lambda x: x["score"], reverse=True)
        top_results = results[:5]  # Prendre plus de contexte pour le LLM
        
        if not top_results or top_results[0]["score"] < 0.05:
            return "❌ Aucun passage pertinent trouvé pour cette question."
        
        # Construire le contexte pour le LLM
        context_parts = []
        source_files = []
        
        for result in top_results:
            context_parts.append(result["text"])
            doc_info = index["documents"].get(result["doc_hash"], {})
            if doc_info.get("filename"):
                source_files.append(doc_info["filename"])
        
        context = "\n\n---\n\n".join(context_parts)
        
        # Générer une réponse intelligente avec Gemini
        llm_response = _generate_response_with_llm(question, context, source_files)
        
        if llm_response:
            return llm_response
        
        # Fallback: utiliser le mode simple si LLM échoue (quota dépassé, etc.)
        return _simple_answer_from_context(question, context, source_files)
        
    except Exception as e:
        logger.error(f"Erreur réponse: {e}")
        return f"❌ Erreur: {str(e)}"


# ═══════════════════════════════════════════════════════════════════════════════
# FONCTIONS DIRECTES (pour appel depuis app.py - sans décorateur @tool)
# ═══════════════════════════════════════════════════════════════════════════════

def process_document_direct(filepath: str) -> str:
    """
    Version directe de process_uploaded_document pour appel depuis l'interface.
    Traite un document uploadé : extraction, chunking et indexation.
    """
    try:
        logger.info(f"📄 Traitement document: {filepath}")
        
        if not os.path.exists(filepath):
            return f"❌ Fichier non trouvé: {filepath}"
        
        filename = os.path.basename(filepath)
        file_hash = get_file_hash(filepath)
        
        # Extraire le contenu
        content = extract_document_content(filepath)
        
        if content.startswith("Erreur"):
            return content
        
        # Découper en chunks
        chunks = chunk_text(content)
        
        if not chunks:
            return "❌ Aucun contenu extractible du document"
        
        # Créer les embeddings et indexer
        index = load_embeddings_index()
        
        # Enregistrer le document
        doc_info = {
            "filename": filename,
            "filepath": filepath,
            "hash": file_hash,
            "processed_at": datetime.now().isoformat(),
            "total_chunks": len(chunks),
            "total_chars": len(content)
        }
        index["documents"][file_hash] = doc_info
        
        # Supprimer les anciens chunks de ce document (si re-upload)
        index["chunks"] = [c for c in index["chunks"] if c.get("doc_hash") != file_hash]
        
        # Indexer les nouveaux chunks
        for i, chunk in enumerate(chunks):
            chunk_entry = {
                "doc_hash": file_hash,
                "chunk_id": i,
                "text": chunk,
                "embedding": simple_embedding(chunk)
            }
            index["chunks"].append(chunk_entry)
        
        save_embeddings_index(index)
        
        logger.info(f"✅ Document indexé: {len(chunks)} chunks")
        
        return f"""✅ Document traité avec succès !

📄 Fichier: {filename}
📊 Statistiques:
- {len(chunks)} sections indexées
- {len(content):,} caractères extraits
- Prêt pour les questions !

💡 Vous pouvez maintenant poser des questions sur ce document."""
        
    except Exception as e:
        logger.error(f"Erreur traitement document: {e}")
        return f"❌ Erreur traitement: {str(e)}"


def search_documents_direct(query: str, top_k: int = 5) -> str:
    """Version directe de search_in_documents pour appel depuis l'interface."""
    try:
        logger.info(f"🔍 Recherche: {query[:50]}...")
        
        index = load_embeddings_index()
        
        if not index["chunks"]:
            return "❌ Aucun document indexé. Veuillez d'abord uploader un document."
        
        # Créer l'embedding de la requête
        query_embedding = simple_embedding(query)
        
        # Calculer les similarités
        results = []
        for chunk in index["chunks"]:
            similarity = cosine_similarity(query_embedding, chunk["embedding"])
            
            # Bonus si les mots de la requête sont présents
            query_words = set(re.findall(r'\b\w+\b', query.lower()))
            chunk_words = set(re.findall(r'\b\w+\b', chunk["text"].lower()))
            word_overlap = len(query_words & chunk_words) / max(len(query_words), 1)
            
            score = similarity * 0.4 + word_overlap * 0.6
            
            results.append({
                "score": score,
                "text": chunk["text"],
                "doc_hash": chunk["doc_hash"],
                "chunk_id": chunk["chunk_id"]
            })
        
        # Trier par score
        results.sort(key=lambda x: x["score"], reverse=True)
        top_results = results[:top_k]
        
        if not top_results or top_results[0]["score"] < 0.1:
            return "❌ Aucun passage pertinent trouvé pour cette question."
        
        # Formater les résultats
        output = ["📚 RÉSULTATS DE RECHERCHE", "=" * 50]
        
        for i, result in enumerate(top_results, 1):
            doc_info = index["documents"].get(result["doc_hash"], {})
            filename = doc_info.get("filename", "Document inconnu")
            
            output.append(f"\n📄 Résultat {i} (score: {result['score']:.2f})")
            output.append(f"Source: {filename}")
            output.append("-" * 40)
            
            text = result["text"]
            if len(text) > 800:
                text = text[:800] + "..."
            output.append(text)
        
        return "\n".join(output)
        
    except Exception as e:
        logger.error(f"Erreur recherche: {e}")
        return f"❌ Erreur recherche: {str(e)}"


def list_documents_direct() -> str:
    """Version directe de list_documents pour appel depuis l'interface."""
    try:
        index = load_embeddings_index()
        
        if not index["documents"]:
            return "📭 Aucun document indexé pour le moment.\n💡 Uploadez un document pour commencer !"
        
        output = ["📚 DOCUMENTS INDEXÉS", "=" * 50]
        
        for i, (doc_hash, doc_info) in enumerate(index["documents"].items(), 1):
            filename = doc_info["filename"]
            chunks = doc_info["total_chunks"]
            chars = doc_info["total_chars"]
            date = doc_info.get("processed_at", "")[:10]
            
            output.append(f"\n{i}. 📄 {filename}")
            output.append(f"   └─ {chunks} sections | {chars:,} caractères | {date}")
        
        output.append(f"\n{'=' * 50}")
        output.append(f"📊 Total: {len(index['documents'])} document(s)")
        
        return "\n".join(output)
        
    except Exception as e:
        logger.error(f"Erreur liste: {e}")
        return f"❌ Erreur: {str(e)}"


def has_indexed_documents() -> bool:
    """Vérifie s'il y a des documents indexés."""
    try:
        index = load_embeddings_index()
        return len(index.get("documents", {})) > 0
    except:
        return False


def get_indexed_filenames() -> List[str]:
    """Retourne la liste des noms de fichiers indexés."""
    try:
        index = load_embeddings_index()
        return [doc["filename"] for doc in index.get("documents", {}).values()]
    except:
        return []


def clear_all_documents() -> str:
    """
    Efface tous les documents indexés et les fichiers uploadés.
    Utilisé pour nettoyer au démarrage de chaque session.
    """
    try:
        import shutil
        
        # Supprimer l'index des embeddings
        if os.path.exists(EMBEDDINGS_FILE):
            os.remove(EMBEDDINGS_FILE)
            logger.info("🗑️ Index des embeddings supprimé")
        
        # Supprimer tous les fichiers dans docDB (sauf le dossier lui-même)
        if os.path.exists(DOC_DB_PATH):
            for filename in os.listdir(DOC_DB_PATH):
                file_path = os.path.join(DOC_DB_PATH, filename)
                try:
                    if os.path.isfile(file_path):
                        os.remove(file_path)
                        logger.info(f"🗑️ Fichier supprimé: {filename}")
                except Exception as e:
                    logger.error(f"Erreur suppression {filename}: {e}")
        
        logger.info("✅ Tous les documents ont été effacés (nouvelle session)")
        return "✅ Documents de la session précédente effacés."
        
    except Exception as e:
        logger.error(f"Erreur nettoyage: {e}")
        return f"❌ Erreur nettoyage: {str(e)}"


# Export des tools (pour CrewAI) et fonctions directes (pour app.py)
__all__ = [
    # Tools CrewAI (avec décorateur @tool)
    'process_uploaded_document',
    'search_in_documents', 
    'summarize_document',
    'list_documents',
    'answer_from_document',
    # Fonctions directes (sans décorateur)
    'process_document_direct',
    'search_documents_direct',
    'list_documents_direct',
    'has_indexed_documents',
    'get_indexed_filenames',
    'clear_all_documents'
]
